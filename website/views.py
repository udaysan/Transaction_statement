from flask import Blueprint,render_template, request
from .models import db, Customer, Account,Transaction, Card, Statement_request
from .import db
from flask_login import  login_required,  current_user
from datetime import datetime
from sqlalchemy import and_, between,desc
from docx import Document
from website.mail_pdf import mail_pdf
import os

views=Blueprint('views',__name__)

@views.route('/')
@login_required
def home():
    return render_template("home.html",customer=current_user)

@views.route('/statements',methods=['GET','POST'])
@login_required
def statements():
    if request.method=='GET':
        return render_template("statement.html",customer=current_user)
    
    if request.method=='POST':
        substring=request.form.get('templateID')
        email_address=request.form.get('emailadd')
        from_date=request.form.get('FromDate')
        to_date=request.form.get('ToDate')
        file_password=request.form.get('password')
        request_datetime=datetime.today()
        user_email_id=current_user.email
        from_date=datetime.strptime(from_date, '%Y-%m-%d').date()
        to_date=datetime.strptime(to_date, '%Y-%m-%d').date()
        new_statement_request = Statement_request(email_address=email_address,from_date=from_date,to_date=to_date,request_datetime=request_datetime,file_password=file_password,user_email_id=user_email_id)
        db.session.add(new_statement_request)
        db.session.commit()
        customer = Customer.query.filter_by(email=user_email_id).first()
        account=Account.query.filter_by(customer_id=customer.id).first()
        sum_debit=0
        sum_credit=0
        sum_total=0
        debit_card_id=""
        credit_card_id=""
        current_directory = os.getcwd()

        if account:

            if substring=="net banking":
                mail_subject='Net Banking Statement PDF'
                relative_path = f'website/static/word_file/net_banking_template.docx'
                template_path = os.path.join(current_directory, relative_path) 
            elif substring=='debit':
                mail_subject='Debit Card Statement PDF'
                card_id_d=Card.query.filter(and_(Card.customer_id==Customer.id,Card.card_type=="Debit")).first()
                debit_card_id=card_id_d.card_number
                relative_path = f'website/static/word_file/debit_card_template.docx'
                template_path = os.path.join(current_directory, relative_path)                 
            elif substring=='credit':
                mail_subject='Credit Card Statement PDF'
                card_id=Card.query.filter(and_(Card.customer_id==Customer.id,Card.card_type=="Credit")).first()
                credit_card_id=card_id.card_number
                relative_path = f'website/static/word_file/credit_card_template.docx'
                template_path = os.path.join(current_directory, relative_path)                  
            elif substring=='upi':
                mail_subject='UPI Transactions Statement PDF'
                relative_path = f'website/static/word_file/upi_template.docx'
                template_path = os.path.join(current_directory, relative_path) 
            else:
                mail_subject='Account Statement PDF'
                relative_path = f'website/static/word_file/all_transactions_template.docx'
                template_path = os.path.join(current_directory, relative_path) 
            
            transactions_total=Transaction.query.filter(and_(Transaction.account_id==account.id)).all()
            transactions = Transaction.query.filter(and_(Transaction.account_id == account.id,
            Transaction.description.ilike(f'%{substring}%'),
            between(Transaction.transaction_datetime, from_date, to_date))).order_by(desc(Transaction.id)).all()
            # Load your pre-designed Word template
            # Replace with your template file path
            doc = Document(template_path)
            column_names_table1=['ID','Description','Date-time','Credit','Debit']
            for table in doc.tables:
                # Check if the table has the same number of columns as target_column_names
                if len(table.rows[0].cells) == len(column_names_table1):
                    table1=table

            for i in transactions_total:
                sum_total=sum_total+i.transaction_amount 

            for transaction in transactions:
                # Customize this part to populate data from your Table2 model
                row = table1.add_row().cells
                transaction.transaction_datetime=transaction.transaction_datetime.strftime("%Y-%m-%d_%H:%M") 
                row[0].text = str(transaction.id)
                row[1].text = str(transaction.description)
                row[2].text = str(transaction.transaction_datetime)
                if transaction.transaction_amount>0:       
                    row[3].text=str(transaction.transaction_amount)
                    row[4].text=""
                    sum_credit=sum_credit+transaction.transaction_amount
                else:
                    row[3].text=""
                    row[4].text = str(abs(transaction.transaction_amount))
                    sum_debit=sum_debit+abs(transaction.transaction_amount)
            
            date_on_doc=str(datetime.now())
            sum_debit = round(sum_debit, 2)
            sum_credit=round(sum_credit,2)
            

            replacements = {
                "var_customer_name_var": current_user.govname,
                "var_statement_generation_date_var": date_on_doc,
                "var_credit_amount_var": "₹ "+str(sum_credit),
                "var_debit_amount_var":"₹ "+str(sum_debit),
                "var_balance_var":"₹ "+str(sum_total),
                "var_card_number_var": credit_card_id,
                "var_debit_number_var":debit_card_id,
                "var_account_number_var":account.id,
                # Add more replacements as needed
            }

            for table in doc.tables:
                if len(table.rows[0].cells) != len(column_names_table1):
                    for row in table.rows:
                        for cell in row.cells:
                            # Iterate through paragraphs in the cell
                            for paragraph in cell.paragraphs:
                                for old_string, new_string in replacements.items():
                                    # Replace old_string with new_string in the paragraph text
                                    new_string=str(new_string)
                                    paragraph.text = paragraph.text.replace(old_string, new_string)

            current_time = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")  # Format the current datetime
            relative_path_output_word_path = f'website/static/user_statement_files/{current_user.email}_{current_time}.docx'
            output_word_path = os.path.join(current_directory, relative_path_output_word_path) 
            doc.save(output_word_path)
            email_sender=''#enter mail id
            email_password=''#enter device specific mail password

            # Convert the Word document to PDF using ReportLab
            relative_path_output_pdf_path = f'website/static/user_statement_files/{current_user.email}_{current_time}.pdf'
            output_pdf_path = os.path.join(current_directory, relative_path_output_pdf_path)             
            #convert(output_word_path)
            mail_pdf(output_word_path, email_address, email_sender, email_password,file_password,mail_subject)

            print(f"Transactions added to {output_word_path} and converted to {output_pdf_path}")

        else:
            print("User don't have an account in the bank!")

        return render_template("statement.html",customer=current_user)

@views.route('/PersonalLoans',methods=['GET'])
@login_required
def PersonalLoans():
    if request.method=='GET':
        return render_template("personal_loans.html",customer=current_user)
    
@views.route('/HomeLoans',methods=['GET'])
@login_required
def HomeLoans():
    if request.method=='GET':
        return render_template("home_loans.html",customer=current_user)
    
@views.route('/fixedDeposit',methods=['GET'])
@login_required
def fixedDeposit():
    if request.method=='GET':
        return render_template("fixed_deposite.html",customer=current_user)